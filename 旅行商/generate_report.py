import json
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).parent
RESULTS_PATH = BASE_DIR / "tsp_results.json"
CHART_PATH = BASE_DIR / "tsp_distance_comparison.png"
REPORT_PATH = BASE_DIR / "实验二-TSP求解实验报告.docx"


def load_data():
    return json.loads(RESULTS_PATH.read_text(encoding="utf-8"))


def make_chart(results):
    names = ["Held-Karp", "Nearest\nNeighbor", "2-opt", "GA", "ACO"]
    keys = ["Held-Karp动态规划", "最近邻算法", "2-opt局部优化", "遗传算法", "蚁群算法"]
    distances = [results[name]["distance"] for name in keys]

    plt.figure(figsize=(8, 4.5))
    bars = plt.bar(names, distances, color=["#4C78A8", "#F58518", "#54A24B", "#E45756", "#72B7B2"])
    plt.ylabel("Tour distance")
    plt.title("TSP Algorithm Distance Comparison")
    plt.ylim(min(distances) - 3, max(distances) + 5)
    for bar, distance in zip(bars, distances):
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            distance + 0.35,
            f"{distance:.2f}",
            ha="center",
            va="bottom",
            fontsize=9,
        )
    plt.tight_layout()
    plt.savefig(CHART_PATH, dpi=180)
    plt.close()


def set_font(document):
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def add_code(document, code):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def make_report(data):
    results = data["results"]
    optimum = results["Held-Karp动态规划"]["distance"]
    nearest = results["最近邻算法"]
    local = results["2-opt局部优化"]
    ga = results["遗传算法"]
    aco = results["蚁群算法"]
    reduced = nearest["distance"] - local["distance"]

    document = Document()
    set_font(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("南京信息工程大学人工智能实验（实习）报告")
    title_run.bold = True
    title_run.font.size = Pt(16)

    info = document.add_table(rows=4, cols=4)
    info.style = "Table Grid"
    info_data = [
        ("实验名称", "旅行商问题（TSP）求解", "日期", "2026.05.25"),
        ("指导教师", "薛羽", "系", "计算机学院、软件学院"),
        ("专业", "软件工程", "年级班次", "2024级 4班"),
        ("姓名", "胡永平", "学号", "202413440119"),
    ]
    for row, values in zip(info.rows, info_data):
        for cell, value in zip(row.cells, values):
            cell.text = value

    document.add_heading("1. 实验目的", level=1)
    document.add_paragraph(
        "理解旅行商问题的建模方法和组合优化特征；掌握精确算法与启发式算法在求解质量、"
        "运行时间方面的差异；使用 Python 实现 Held-Karp 动态规划、最近邻算法、"
        "2-opt 局部优化、遗传算法和蚁群算法，并对结果进行比较分析。"
    )

    document.add_heading("2. 实验内容和实验方法", level=1)
    document.add_heading("2.1 实验内容", level=2)
    document.add_paragraph(
        "给定 12 个二维城市坐标，以欧氏距离作为城市间距离，寻找访问每个城市一次并返回"
        "出发城市 0 的最短闭合路径。按照实验要求，不实现模拟退火算法；以 Held-Karp "
        "得到的最优路线作为对照，评价四种启发式方法的求解效果。"
    )
    document.add_heading("2.2 实验方法", level=2)
    methods = [
        "Held-Karp 动态规划：以“已访问城市集合 + 当前终点”为状态，得到小规模问题的精确最优解。",
        "最近邻算法：从城市 0 开始，每步选择最近的未访问城市，快速构造初始可行解。",
        "2-opt 局部优化：反转路线中的连续片段，接受能缩短总距离的交换结果。",
        "遗传算法：使用排列编码、锦标赛选择、顺序交叉、交换变异和精英保留搜索路线。",
        "蚁群算法：蚂蚁根据距离启发信息与信息素选择下一城市，反复更新信息素得到较优路线。",
    ]
    for method in methods:
        document.add_paragraph(method, style="List Bullet")

    document.add_heading("3. 实验步骤和实验结果", level=1)
    document.add_heading("3.1 问题分析与状态定义", level=2)
    document.add_paragraph(
        "一条路线表示为 [0, c1, c2, ..., c11, 0]，其中中间城市不重复。目标函数为相邻"
        "城市间欧氏距离之和。动态规划采用二进制 mask 表示已访问城市集合；遗传算法的染色体"
        "为城市 1 至 11 的排列；蚁群算法在每对城市之间维护信息素。随机算法均采用 seed=42，"
        "保证本报告中的实验数据可以复现。"
    )

    document.add_heading("3.2 代码编写", level=2)
    document.add_paragraph(
        "程序文件为 TSP.py，统一提供距离计算、五种算法与实验入口。以下代码展示实验调度方式，"
        "各算法返回路线和总距离，再统一计算相对于最优解的误差。"
    )
    add_code(
        document,
        'algorithms = [\n'
        '    ("Held-Karp动态规划", lambda: held_karp(distances)),\n'
        '    ("最近邻算法", lambda: nearest_neighbor(distances)),\n'
        '    ("2-opt局部优化", lambda: two_opt(nearest_neighbor(distances)[0], distances)),\n'
        '    ("遗传算法", lambda: genetic_algorithm(distances, seed=42)),\n'
        '    ("蚁群算法", lambda: ant_colony(distances, seed=42)),\n'
        "]"
    )

    document.add_heading("3.3 运行结果", level=2)
    result_table = document.add_table(rows=1, cols=4)
    result_table.style = "Table Grid"
    for cell, text in zip(
        result_table.rows[0].cells, ["算法", "总距离", "相对最优误差/%", "运行时间/ms"]
    ):
        cell.text = text
    for name, result in results.items():
        cells = result_table.add_row().cells
        cells[0].text = name
        cells[1].text = f"{result['distance']:.2f}"
        cells[2].text = f"{result['error_percent']:.2f}"
        cells[3].text = f"{result['time_ms']:.3f}"

    document.add_paragraph()
    document.add_picture(str(CHART_PATH), width=Inches(5.9))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = document.add_paragraph("图 1  不同算法求解所得路线长度比较")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("3.4 路线输出", level=2)
    for name, result in results.items():
        route_text = " -> ".join(str(city) for city in result["route"])
        document.add_paragraph(f"{name}：{route_text}")

    document.add_heading("4. 分析与讨论", level=1)
    document.add_paragraph(
        f"Held-Karp 动态规划给出的最优总距离为 {optimum:.2f}，因此可作为本次小规模实验的"
        "评价基准。最近邻算法运行最快，但其路线距离为 "
        f"{nearest['distance']:.2f}，比最优解高 {nearest['error_percent']:.2f}%，说明局部"
        "贪心选择并不能保证整体路线最短。"
    )
    document.add_paragraph(
        f"2-opt 在最近邻路线基础上将距离减少 {reduced:.2f}，误差由 "
        f"{nearest['error_percent']:.2f}% 降到 {local['error_percent']:.2f}%。这表明简单的"
        "局部交换即可明显修正最近邻算法产生的不合理边，但仍可能停在局部最优位置。"
    )
    document.add_paragraph(
        f"遗传算法和蚁群算法本次分别获得 {ga['distance']:.2f} 与 {aco['distance']:.2f} "
        "的路线，与动态规划基准一致。二者通过多次搜索提升了解质量，但运行时间分别为 "
        f"{ga['time_ms']:.3f} ms 和 {aco['time_ms']:.3f} ms，明显高于最近邻与 2-opt。"
        "因此，当城市规模较小且需要确认最优性时可采用动态规划；当规模扩大时，可用 "
        "2-opt 获得快速改进结果，或采用遗传算法、蚁群算法换取更高质量的近似解。"
    )

    document.add_heading("5. 实验结论", level=1)
    document.add_paragraph(
        "本实验完成了 TSP 的精确求解与四种启发式求解算法实现。结果显示，最近邻适合快速"
        "构造路线，2-opt 能以很小计算代价提高路线质量，遗传算法和蚁群算法在本实验数据上"
        "获得了最优距离。实验同时说明，算法选择需要在求解质量和计算开销之间进行权衡。"
    )

    document.save(REPORT_PATH)


def main():
    data = load_data()
    make_chart(data["results"])
    make_report(data)
    print(CHART_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
